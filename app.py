import streamlit as st
import openai
from openai import OpenAI
import os
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

# Load environment variables
load_dotenv()

# Initialize OpenAI client
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# Page configuration
st.set_page_config(
    page_title="AI Resume Builder",
    page_icon="📄",
    layout="wide"
)

# Custom CSS for better UI
st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .stButton>button {
        width: 100%;
        background-color: #4CAF50;
        color: white;
        font-weight: bold;
        padding: 0.5rem 1rem;
        border-radius: 0.5rem;
    }
    .stButton>button:hover {
        background-color: #45a049;
    }
    h1 {
        color: #2c3e50;
        text-align: center;
    }
    .success-box {
        padding: 1rem;
        background-color: #d4edda;
        border-left: 4px solid #28a745;
        border-radius: 0.25rem;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

def generate_resume_content(original_resume_text, position_description):
    """
    Use OpenAI to generate an optimized resume based on the original resume and position description.
    """
    try:
        prompt = f"""You are an expert resume writer and career consultant. Based on the following information, create a professional, ATS-friendly resume that is tailored to the specific position.

ORIGINAL RESUME:
{original_resume_text}

TARGET POSITION DESCRIPTION:
{position_description}

Please create an optimized resume that:
1. Highlights relevant skills and experiences from the original resume that match the position
2. Uses keywords from the position description naturally
3. Maintains truthfulness - only use information from the original resume
4. Follows a professional format with clear sections
5. Quantifies achievements where possible
6. Is concise and impactful

Format the resume with these sections (use only sections that are relevant):
- PROFESSIONAL SUMMARY
- SKILLS
- WORK EXPERIENCE
- EDUCATION
- CERTIFICATIONS (if applicable)
- PROJECTS (if applicable)

Return the resume in a clean, well-formatted text structure."""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are an expert resume writer who creates ATS-friendly, professional resumes tailored to specific job positions."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        
        return response.choices[0].message.content
    
    except Exception as e:
        st.error(f"Error generating resume: {str(e)}")
        return None

def create_word_document(resume_content, applicant_name):
    """
    Create a professionally formatted Word document from the resume content.
    """
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Split content into lines
    lines = resume_content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        
        # Check if line is a section header (all caps or ends with colon)
        if line.isupper() or (line.endswith(':') and len(line) < 50):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0, 0, 0)
            p.space_after = Pt(6)
        # Check if line starts with a bullet or dash
        elif line.startswith(('•', '-', '*')):
            p = doc.add_paragraph(line.lstrip('•-* '), style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.runs[0]
            run.font.size = Pt(11)
        else:
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.font.size = Pt(11)
    
    # Save to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    return doc_io

def main():
    # Header
    st.title("📄 AI-Powered Resume Builder")
    st.markdown("### Transform your resume to match any job position using AI")
    
    st.markdown("---")
    
    # Create two columns for layout
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("📋 Your Information")
        
        # Applicant name
        applicant_name = st.text_input("Your Full Name", placeholder="John Doe")
        
        # Resume upload
        uploaded_file = st.file_uploader(
            "Upload Your Current Resume",
            type=['txt', 'pdf', 'docx'],
            help="Upload your existing resume in TXT, PDF, or DOCX format"
        )
        
        # Manual text input option
        st.markdown("**OR paste your resume text below:**")
        resume_text = st.text_area(
            "Resume Content",
            height=300,
            placeholder="Paste your resume content here..."
        )
    
    with col2:
        st.subheader("🎯 Target Position")
        
        # Position/Job description
        position_description = st.text_area(
            "Job Description / Position Requirements",
            height=400,
            placeholder="Paste the job description or position requirements here...\n\nInclude:\n- Job title\n- Required skills\n- Responsibilities\n- Qualifications\n- Preferred experience"
        )
    
    st.markdown("---")
    
    # Process uploaded file
    original_resume_text = resume_text
    
    if uploaded_file is not None:
        try:
            if uploaded_file.type == "text/plain":
                original_resume_text = uploaded_file.read().decode("utf-8")
            elif uploaded_file.type == "application/pdf":
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                original_resume_text = ""
                for page in pdf_reader.pages:
                    original_resume_text += page.extract_text()
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                doc = Document(uploaded_file)
                original_resume_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            st.error(f"Error reading file: {str(e)}")
    
    # Generate button
    col_btn1, col_btn2, col_btn3 = st.columns([1, 2, 1])
    with col_btn2:
        generate_button = st.button("🚀 Generate Optimized Resume", use_container_width=True)
    
    # Generate resume
    if generate_button:
        if not original_resume_text:
            st.error("⚠️ Please provide your resume content either by uploading a file or pasting text.")
        elif not position_description:
            st.error("⚠️ Please provide the target position description.")
        elif not applicant_name:
            st.error("⚠️ Please enter your full name.")
        else:
            with st.spinner("🤖 AI is crafting your optimized resume..."):
                generated_resume = generate_resume_content(original_resume_text, position_description)
                
                if generated_resume:
                    st.success("✅ Resume generated successfully!")
                    
                    # Store in session state
                    st.session_state['generated_resume'] = generated_resume
                    st.session_state['applicant_name'] = applicant_name
    
    # Display generated resume
    if 'generated_resume' in st.session_state:
        st.markdown("---")
        st.subheader("📝 Your Optimized Resume")
        
        # Display resume in a nice box
        st.markdown(f"""
        <div style="background-color: #f8f9fa; padding: 2rem; border-radius: 0.5rem; border: 1px solid #dee2e6;">
            <pre style="white-space: pre-wrap; font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; font-size: 14px; line-height: 1.6;">
{st.session_state['generated_resume']}
            </pre>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Download options
        st.subheader("💾 Download Your Resume")
        
        col_dl1, col_dl2 = st.columns(2)
        
        with col_dl1:
            # Download as TXT
            st.download_button(
                label="📄 Download as TXT",
                data=st.session_state['generated_resume'],
                file_name=f"{st.session_state['applicant_name'].replace(' ', '_')}_Resume.txt",
                mime="text/plain",
                use_container_width=True
            )
        
        with col_dl2:
            # Download as DOCX
            try:
                doc_file = create_word_document(
                    st.session_state['generated_resume'],
                    st.session_state['applicant_name']
                )
                st.download_button(
                    label="📝 Download as DOCX",
                    data=doc_file,
                    file_name=f"{st.session_state['applicant_name'].replace(' ', '_')}_Resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            except Exception as e:
                st.error(f"Error creating Word document: {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown("""
        <div style="text-align: center; color: #6c757d; padding: 1rem;">
            <p>💡 <strong>Tip:</strong> Review and customize the generated resume to ensure it accurately represents your experience.</p>
            <p style="font-size: 0.9rem;">Powered by OpenAI GPT-4</p>
        </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
